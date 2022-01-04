"""
Genuary2022 Day 1 - Draw 1000 of something
C44 one liner inspired Maze by C. Ponsard
Free under the terms of GPLv3 license

Will generate as word file
Tune paragraph spacing manually - TODO improve generator

Requires
- python3.x
- python-docx
- C64 Pro Mono font (for generated word)
"""

from docx import Document
from docx.shared import Pt
import random

document = Document()
t=["",""]

for l in range(100):
    rs = ""
    for i in range(100):
        rs += t[random.randrange(0, 2)]
    p = document.add_paragraph()
    run = p.add_run(rs)
    run.font.name = 'C64 Pro Mono'
    run.font.size = Pt(4)

document.add_page_break()
document.save('day1_maze.docx')